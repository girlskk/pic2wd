from copy import deepcopy
from datetime import datetime
import os

from docx import Document
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from concurrent.futures import ThreadPoolExecutor, wait, FIRST_EXCEPTION
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches
from inputimeout import inputimeout, TimeoutOccurred


bugMap = {
    "杆塔树障": "基础",
    "杆塔未封顶": "基础",
    "杆塔异物": "基础",
    "施工遗留": "基础",
    "杆塔鸟巢": "基础",
    "杆塔倾斜": "基础",
    "塔基植被覆盖": "基础",
    "塔基杂物堆积": "基础",
    "杆塔倾斜": "基础",
    "塔基树障": "基础",
    "杆塔异物": "基础",
    "杆塔裂纹": "基础",
    "杆塔未封顶": "基础",
    "杆塔损伤": "基础",
    "塔头破损": "基础",
    "杆塔破损": "基础",
    "拉线松弛": "基础",
    "横担锈蚀": "基础",
    "绝缘子脱落": "绝缘子",
    "绝缘子破损": "绝缘子",
    "绝缘子老化": "绝缘子",
    "绝缘子倾斜": "绝缘子",
    "绝缘子污秽": "绝缘子",
    "绝缘子灼伤": "绝缘子",
    "绝缘子雷击": "绝缘子",
    "釉面剥落": "绝缘子",
    "绑带松脱": "绝缘子",
    "绝缘子绑带安装不规范": "绝缘子",
    "金具锈蚀": "金具",
    "销钉缺失": "金具",
    "销钉退出": "金具",
    "销钉安装不规范": "金具",
    "螺母松动": "金具",
    "螺母缺失": "金具",
    "防震锤锈蚀": "金具",
    "防震锤脱落": "金具",
    "导线缠绕": "导地线",
    "导线脱落": "导地线",
    "导线悬挂异物": "导地线",
    "导线断股": "导地线",
    "导线松股": "导地线",
    "导线固定不牢": "导地线",
    "地线悬挂异物": "导地线",
    "绝缘保护壳破损": "附属设施",
    "绝缘保护壳缺失": "附属设施",
    "标识牌脱落": "附属设施",
    "通道树障": "通道",
    "通道施工": "通道",
    "变压器漏油": "变压器",
    "变压器渗油": "变压器",
    "避雷器雷击": "避雷器",
    "避雷器破损": "避雷器",
    "线耳脱落": "避雷器",
    "避雷器连接线脱落": "避雷器",
}

bug_type_count_map = {}
total_statis_map = {}
image_index = {}
bug_type_map = {1: "危急", 2: "严重", 3: "一般"}
close_up_map = {}
image_bug_level_map = {}
image_bug_reason_map = {}
image_tower_map = {}
image_route_name_map = {}
image_type_map = {}
pic_name_cache = {}


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


# 缺陷描述和缺陷类别不匹配时，模糊匹配
def fuzzy_match(bug_detail=""):
    """When the exact match fails, fuzzy matching is used"""
    if any(keyword in bug_detail for keyword in ["绝缘子"]):
        return "绝缘子"
    if any(keyword in bug_detail for keyword in ["杆塔", "塔基", "塔头", "塔顶"]):
        return "基础"
    if any(keyword in bug_detail for keyword in ["金具", "销钉", "螺母"]):
        return "金具"
    if any(keyword in bug_detail for keyword in ["保护壳", "标识牌"]):
        return "附属设施"
    if any(keyword in bug_detail for keyword in ["地线", "导线"]):
        return "导地线"
    if any(keyword in bug_detail for keyword in ["避雷器"]):
        return "避雷器"
    if any(keyword in bug_detail for keyword in ["变压器"]):
        return "变压器"
    if any(keyword in bug_detail for keyword in ["通道"]):
        return "通道"
    debug_log(f"{bug_detail} 未匹配到缺陷类别", 2)
    return ""


def get_bug_type(bug_reason=""):
    """Get the defect type based on the defect description"""
    bugType = bugMap.get(bug_reason, "")
    if len(bugType) == 0:
        bugType = fuzzy_match(bug_reason)
        if len(bugType) > 0:
            debug_log(
                f"缺陷描述:\033[32m[{bug_reason}]\033[m 未匹配到缺陷类别,已模糊匹配为 >>> \033[32m{bugType}\033[m",
                1,
            )
    return bugType


# statis_add_table 汇总数据写入
def set_detail_statis(table, images, bug_type):
    debug_log(f"开始处理 {bug_type_map.get(bug_type,'')}缺陷汇总表")
    c = 1
    for i in images:
        picName = get_pic_name(i)
        bugLevel = image_bug_level_map.get(i, "")

        bugType = get_bug_type(image_bug_reason_map.get(i, ""))
        # 汇总数据
        if len(bugType) > 0:
            update_bug_type_count(bugType, bugLevel)

        update_cell(table, c, 0, str(c))
        update_cell(table, c, 1, picName[:-3])
        update_cell(table, c, 2, bugType)
        update_cell(table, c, 3, bugLevel)
        update_cell(table, c, 4, image_index.get(i, ""))
        c += 1
    debug_log(f"{bug_type_map.get(bug_type,'')}缺陷汇总表 写入完成")


# 更新单元格
def update_cell(table, row_idx, col_idx, text):
    """Update the cell and set it to center"""
    cell = table.cell(row_idx, col_idx)
    cell.text = text
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# 单元格宽高
def set_cell_size(table, row_idx, col_idx, width, height):
    cell = table.cell(row_idx, col_idx)
    cell.width = Cm(width)
    cell.height = Cm(height)


# 设置单元格居中
def cell_set_center(cell):
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# 复制单元格字体大小
def copy_cell_font_size(cell1, cell2):
    cell1.paragraphs[0].runs[0].font.size = cell2.paragraphs[0].runs[0].font.size


def update_bug_type_count(bugType, bugLevel):
    bug_type_count_map.setdefault(bugType, {}).update(
        {
            bugLevel: bug_type_count_map.get(bugType, {}).get(bugLevel, 0) + 1,
            "合计": bug_type_count_map.get(bugType, {}).get("合计", 0) + 1,
        }
    )
    bug_type_count_map.setdefault("合计", {}).update(
        {
            bugLevel: bug_type_count_map.get("合计", {}).get(bugLevel, 0) + 1,
            "合计": bug_type_count_map.get("合计", {}).get("合计", 0) + 1,
        }
    )


def get_pic_name(pic):
    cache_name = pic_name_cache.get(pic, "")
    if len(cache_name) > 0:
        return cache_name
    picName, picType = pic.split(".")
    pic_name_cache[pic] = picName
    return picName


# 每个图片插入数据到一个表格
def deal_table(table, pic):
    debug_log(f"开始处理:{pic}")
    pic_name = get_pic_name(pic)
    route_name, tower_num, bug_reason, bug_level = pic_name.split("_")
    update_cell(table, 1, 0, route_name)
    update_cell(table, 1, 1, tower_num)
    update_cell(table, 1, 2, bug_level)
    update_cell(table, 2, 1, bug_reason)

    insert_image_designation(table, 9, pic, 16.4, 12.3, WD_PARAGRAPH_ALIGNMENT.CENTER)
    close_up_pic = close_up_map.get(pic_name, "")
    table.rows[4].height = Cm(7.34)
    if len(close_up_pic) > 0:
        insert_image_designation(
            table, 12, close_up_pic, 7, 7, WD_PARAGRAPH_ALIGNMENT.LEFT
        )
    debug_log(f"明细表 {pic_name} 处理完成")


#
# 插入图片
def insert_image_by_rate(table, cell_idx, pic, xRate, yRate, alignment):
    image_path = "./pic/" + pic
    cell = table._cells[cell_idx]

    cell.paragraphs[0].add_run().add_picture(
        image_path,
        width=cell.width * xRate,
        height=table.rows[3].height * yRate,
    )
    cell.paragraphs[0].alignment = alignment


def insert_image_designation(table, cell_idx, pic, x, y, alignment):
    image_path = "./pic/" + pic
    cell = table._cells[cell_idx]

    cell.paragraphs[0].add_run().add_picture(
        image_path,
        width=Cm(x),
        height=Cm(y),
    )
    cell.paragraphs[0].alignment = alignment


# 缺陷数量统计表
def bug_num_statis(table):
    bugLevelCountMap = bug_type_count_map.get("合计", {})
    if len(bugLevelCountMap) == 0:
        return
    col = 0
    for key in table.rows[0].cells:
        count = bugLevelCountMap.get(key.text, 0)
        if count > 0:
            table.cell(1, col).text = str(count)
            table.cell(1, col).paragraphs[0].runs[0].font.name = statis_number_font
            if is_set_statis_number_size:
                copy_cell_font_size(table.cell(1, col), table.cell(0, col))
            cell_set_center(table.cell(1, col))
            total_statis_map[key.text] = count
        col += 1


# 缺陷类别统计表
def bug_type_statis(table):
    row = 0
    for rows in table.rows:
        bugLevelCountMap = bug_type_count_map.get(rows.cells[0].text, {})
        if len(bugLevelCountMap) > 0 or row >= 2:
            col = 0
            for key in table.rows[1].cells:
                count = bugLevelCountMap.get(key.text, 0)
                if count > 0 or col >= 1:
                    table.cell(row, col).text = str(count)
                    table.cell(row, col).paragraphs[0].runs[
                        0
                    ].font.name = statis_number_font
                    if is_set_statis_number_size:
                        copy_cell_font_size(table.cell(row, col), table.cell(1, col))

                    cell_set_center(table.cell(row, col))
                    if key.text == "合计":
                        total_statis_map[rows.cells[0].text] = count
                col += 1
        row += 1


# 缺陷情况总览
def set_total_description(doc):
    total_description_paragraph = match_text_paragraph(doc, "本次现场巡检")
    if total_description_paragraph == 0:
        debug_log(" 定位缺陷情况总览模块失败......", 2)
        return False
    para = doc.paragraphs[total_description_paragraph]
    tpl = para.text
    font_size = para.runs[0].font.size
    content = tpl.format(
        total_bug=total_statis_map.get("合计", 0),
        weiji_bug=total_statis_map.get("危急", 0),
        yanzhong_bug=total_statis_map.get("严重", 0),
        yiban_bug=total_statis_map.get("一般", 0),
        bileiqi_bug=total_statis_map.get("避雷器", 0),
        bianyaqi_bug=total_statis_map.get("变压器", 0),
        daodixian_bug=total_statis_map.get("导地线", 0),
        fushu_bug=total_statis_map.get("附属设施", 0),
        jichu_bug=total_statis_map.get("基础", 0),
        jinjv_bug=total_statis_map.get("金具", 0),
        jueyuanzi_bug=total_statis_map.get("绝缘子", 0),
        tongdao_bug=total_statis_map.get("通道", 0),
    )
    debug_log(f"缺陷情况总览文字: {content}")
    doc.paragraphs[total_description_paragraph].text = content
    doc.paragraphs[total_description_paragraph].runs[0].font.size = font_size
    return True


def deal_close_up_image(pic):
    """ "Check whether the close-up image is standard"""
    pic_name, pic_type = pic.split(".")
    if len(pic_name.split("_")) != 5:
        debug_log(
            f"图片名称不规范,不规范的图片为：\033[35m{pic_name}.{pic_type}\033[m ", 2
        )
        return False
    if len(pic_name.split("_特写")) != 2:
        debug_log(
            f"图片名称不规范,不规范的图片为：\033[35m{pic_name}.{pic_type}\033[m ", 2
        )
        return False
    close_up_name, _ = pic_name.split("_特写")
    close_up_map[close_up_name] = pic
    return True


# 获取待处理的图片
def get_images(image_dir=""):
    """Image Classification"""
    image_list = []
    for root, dirs, pics in os.walk(image_dir):
        for pic in pics:
            picName, picType = pic.split(".")
            pic_name_cache[pic] = picName
            if len(picName.split("_")) != 4:
                if not deal_close_up_image(pic):
                    return
                continue
            route_name, tower_num, bug_reason, bug_level = picName.split("_")
            image_bug_level_map[pic] = bug_level
            image_bug_reason_map[pic] = bug_reason
            image_tower_map[pic] = tower_num
            image_route_name_map[pic] = route_name
            image_type_map[pic] = picType
            image_list.append(pic)

    common_list = []
    critical_list = []
    emergency_list = []
    clear_exif(image_list)
    for i in image_list:
        bug_level = image_bug_level_map.get(i, "")
        match bug_level:
            case "危急":
                emergency_list.append(i)
            case "严重":
                critical_list.append(i)
            case "一般":
                common_list.append(i)

    return emergency_list, critical_list, common_list


def get_summary_table_index(doc, index=1):
    """Get summary table location"""
    i = 0
    sort = 1
    for t in doc.tables:
        if t._cells[1].text == "缺陷描述":
            if index == sort:
                return i
            sort += 1
        i += 1


def match_text_paragraph(doc, text=""):
    """Get the paragraph where the text is located"""
    pi = 0
    for p in doc.paragraphs:
        if text in p.text:
            return pi
        pi += 1
    return 0


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def missing_table_num(doc, table_index=0, image_num=0):
    """Get the number of missing detail lists"""
    for i in range(table_index, table_index + image_num):
        if i >= len(doc.tables):
            return table_index + image_num - i
        table = doc.tables[i]
        if i > table_index and (table._cells[0].text != "线路名称"):
            return table_index + image_num - i
    return 0


def get_detail_table_index(doc, index=1):
    """Get the location of the details table"""
    i = 0
    sort = 1
    for t in doc.tables:
        if t._cells[0].text == "线路名称":
            if index == sort:
                return i
            sort += 1
        i += 1


def add_missing_table(doc, tbl=None, paragraph_index=0, add_num=0):
    """Add missing table"""

    if tbl == None:
        debug_log("获取明细表表格模板失败", 2)
        return

    if add_num > 0:
        debug_log(f"危急部分缺少{add_num}个表格")
        for i in range(add_num):
            new_tbl = deepcopy(tbl)
            insert_paragraph_after(
                doc.paragraphs[paragraph_index + i],
            ).add_run().add_break(WD_BREAK.PAGE)
            doc.paragraphs[paragraph_index + i]._p.addnext(new_tbl)
        debug_log(f"<危急>部分插入{add_num}个表格成功")


def add_missing_rows(doc, table_index=0, image_num=0, bug_type=0):
    statis_table_rows = len(doc.tables[table_index].rows) - 1
    if statis_table_rows < image_num:
        table_add_row(doc.tables[table_index], image_num)
        debug_log(f"{bug_type}缺陷汇总表插入{image_num-statis_table_rows-1}行")


# table_add_row 汇总行数小于图片数量时，添加行
def table_add_row(table, num=0):
    """Adding rows to a table"""
    tr = len(table.rows)
    while tr - 1 < num:
        new_row = deepcopy(table.rows[-1])
        table.rows[-1]._tr.addnext(new_row._element)  # 在最后一行后面添加
        table.rows[-1]._tr.addprevious(new_row._element)
        tr += 1


# 生成模板
def get_template(emergencyList, criticalList, commonList, file_name=""):
    # 实例化一个Document对象，相当于打开word软件，新建一个空白文件
    doc = Document(file_name)
    # tables = doc.tables  # 获取文档中所有表格对象的列表

    emergency_detail_table_index = get_detail_table_index(doc, 1)
    critical_detail_table_index = get_detail_table_index(doc, 2)
    common_detail_table_index = get_detail_table_index(doc, 3)
    if not (
        emergency_detail_table_index
        * critical_detail_table_index
        * common_detail_table_index
    ):
        debug_log("定位危急缺陷明细表失败", 2)
        return False
    emergency_statis_table_index = get_summary_table_index(doc, 1)
    critical_statis_table_index = get_summary_table_index(doc, 2)
    common_statis_table_index = get_summary_table_index(doc, 3)

    emergency_detail_paragraph = match_text_paragraph(doc, "危急缺陷明细表")
    if emergency_detail_paragraph == 0:
        debug_log("定位危急缺陷明细表失败", 2)
        return False

    # 判断  汇总表行数是否足够
    add_missing_rows(doc, emergency_statis_table_index, len(emergencyList), "危急")
    add_missing_rows(doc, critical_statis_table_index, len(criticalList), "严重")
    add_missing_rows(doc, common_statis_table_index, len(commonList), "一般")
    tpl_table = doc.tables[emergency_detail_table_index]

    tpl_table.rows[0].height = Cm(0.85)
    tpl_table.rows[1].height = Cm(0.85)
    tpl_table.rows[2].height = Cm(0.85)
    set_cell_size(tpl_table, 0, 0, 8.42, 0.85)
    set_cell_size(tpl_table, 1, 0, 8.42, 0.85)

    set_cell_size(tpl_table, 0, 1, 4.21, 0.85)
    set_cell_size(tpl_table, 1, 1, 4.21, 0.85)

    set_cell_size(tpl_table, 0, 2, 4.21, 0.85)
    set_cell_size(tpl_table, 1, 2, 4.21, 0.85)

    set_cell_size(tpl_table, 2, 0, 8.42, 0.85)
    set_cell_size(tpl_table, 2, 1, 8.42, 0.85)

    tbl = doc.tables[emergency_detail_table_index]._tbl

    # 判断 <危急> 详情表数量是否足够
    addNum = missing_table_num(doc, emergency_detail_table_index, len(emergencyList))
    # <危急> 添加表格
    add_missing_table(doc, tbl, emergency_detail_paragraph, addNum)
    critical_detail_table_index += addNum
    common_detail_table_index += addNum

    # 判断 <严重> 详情表数量是否足够
    critical_detail_paragraph = match_text_paragraph(doc, "严重缺陷明细表")
    if critical_detail_paragraph == 0:
        debug_log("严重缺陷明细表失败", 2)
        return False
    addNum = missing_table_num(doc, critical_detail_table_index, len(criticalList))
    # <严重> 添加表格
    add_missing_table(doc, tbl, critical_detail_paragraph, addNum)
    common_detail_table_index += addNum

    # 判断 <一般> 详情表数量是否足够
    common_detail_paragraph = match_text_paragraph(doc, "一般缺陷明细表")
    if critical_detail_paragraph == 0:
        debug_log("一般缺陷明细表失败", 2)
        return False
    addNum = missing_table_num(doc, common_detail_table_index, len(commonList))
    add_missing_table(doc, tbl, common_detail_paragraph, addNum)

    doc.save("tpl.docx")
    debug_log("生成模板成功")
    return True


def deal_one_type_table(doc, table_index, iamge_list, bug_type):
    debug_log(f"开始处理 {bug_type_map.get(bug_type,'')}明细表")
    bug_pre = ""
    match bug_type:
        case 1:
            bug_pre = "A"
        case 2:
            bug_pre = "B"
        case 3:
            bug_pre = "C"
    if len(iamge_list) == 0:
        insert_row(doc.tables[table_index + 1], 0, [bug_pre + str(1)])
        return
    picIndex = 0
    for i in range(
        table_index + 1,
        table_index + 1 + len(iamge_list),
    ):
        table = doc.tables[i]
        pic = iamge_list[picIndex]
        deal_table(table, pic)
        picIndex += 1
        bug_pre_index = bug_pre + str(picIndex)
        insert_row(table, 0, [bug_pre_index + " " + get_pic_name(pic)])
        # todo
        image_index[pic] = bug_pre_index
    debug_log(f"{bug_type_map.get(bug_type,'')}明细表 处理完成")


def insert_row(table, row_index, content):
    # 在指定位置插入一行
    new_row = table.add_row().cells
    for i, text in enumerate(content):
        new_row[i].text = text
    # 移动新插入的行到指定位置
    rows = table.rows
    rows[row_index]._element.getparent().insert(
        rows[row_index]._element.getparent().index(rows[row_index]._element),
        rows[-1]._element,
    )
    # 需要合并的行和列
    cell_span = table.rows[row_index].cells[:]
    cell_span[0].merge(cell_span[-1])
    # rows[-1]._element.getparent().remove(rows[-1]._element)

    # set_cell_border(
    #     table._cells[0],
    #     top={"color": "#ffffff", "val": "nil"},
    #     left={"color": "#ffffff", "val": "nil"},
    #     right={"color": "#ffffff", "val": "nil"},
    # )
    table.rows[0].height = Cm(0.85)
    cell_set_center(table._cells[0])
    table._cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # set_cell_size(table, 0, 0, 16.84, 0.85)


# 处理数据
def deal(emergencyList, criticalList, commonList, fileName):
    # 处理数据
    doc = Document("tpl.docx")
    tables = doc.tables  # 获取文档中所有表格对象的列表
    emergency_statis_table_index = get_summary_table_index(doc, 1)
    critical_statis_table_index = get_summary_table_index(doc, 2)
    common_statis_table_index = get_summary_table_index(doc, 3)

    deal_one_type_table(doc, emergency_statis_table_index, emergencyList, EMERGENCY)
    deal_one_type_table(doc, critical_statis_table_index, criticalList, CRITICAL)
    deal_one_type_table(doc, common_statis_table_index, commonList, COMMON)

    set_detail_statis(tables[emergency_statis_table_index], emergencyList, EMERGENCY)
    set_detail_statis(tables[critical_statis_table_index], criticalList, CRITICAL)
    set_detail_statis(tables[common_statis_table_index], commonList, COMMON)

    bug_num_statis(tables[bug_num_table_index])
    debug_log("缺陷数量统计表 写入完成")
    bug_type_statis(tables[bug_type_table_index])
    debug_log("缺陷类别统计表 写入完成")

    if set_total_description(doc):
        debug_log("缺陷情况总览 写入完成")
    debug_log("处理结束，正在保存文件...")
    doc.save(fileName)
    debug_log("文件保存文件成功")


def debug_log(message, log_level=0):
    level_tips = ""
    match log_level:
        case 0:
            if not debug:
                return
            level_tips = "[INFO]   "
        case 1:
            level_tips = "\033[33m[WARNING]\033[m"
        case 2:
            level_tips = "\033[31m[ERROR]\033[m  "
    print(f"{level_tips}{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {message}")


# 处理exif信息
def clear_exif(imageList):

    def clear(image):
        debug_log(f"开始清除{image} 的exif信息")
        f = Image.open(image)  # 你的图片文件
        f.save(image)  # 替换掉你的图片文件
        f.close()
        debug_log(f"清除{image} 的exif信息成功")

    executor = ThreadPoolExecutor(ThreadPoolNum)
    all_tasks = [
        executor.submit(clear, IMAGE_DIR + "\\" + imageList[i])
        for i in range(len(imageList))
    ]
    wait(all_tasks, return_when=FIRST_EXCEPTION)


def timer_input(msg="", default="", time_out=60):
    m = ""
    try:
        # 5秒内未完成输入，则超时
        m = inputimeout(
            prompt=f"请在{time_out}秒钟内输入{msg}",
            timeout=time_out,
        )
    except TimeoutOccurred:
        debug_log(f"您未输入...\n使用默认值<\033[32m{default}\033[m>\n", 1)
        m = default
    return m


template_file_name = "template.docx"  # 模板文件名称
bug_num_table_index = 4  # 缺陷数量表位置
bug_type_table_index = bug_num_table_index + 1  # 缺陷类别表位置
statis_number_font = "Times New Roman"  # 统计表数字字体
is_set_statis_number_size = True
debug = True  # 是否开启提示
warn = True  # 是否开启警告信息
ThreadPoolNum = 10
EMERGENCY = 1
CRITICAL = 2
COMMON = 3
IMAGE_DIR = ".\\pic"


def get_path():
    dir = input("请输入(示例:c:\\pic):")
    if os.path.exists(dir):
        return
    debug_log(f"该文件夹不存在，请重新输入", 1)
    get_path()


def main():
    global IMAGE_DIR
    debug_log(
        f"请确认要处理的文件在\033[32m{os.getcwd()}{IMAGE_DIR[1:]}\033[m目录下",
        1,
    )
    # custom_path = timer_input("\n<是>输入:1\n<否>输入:2\n如需选择其他路径,输入:3")
    # match custom_path:
    #     # case "1":

    #     case "2":
    #         return
    #     case "3":
    #         IMAGE_DIR = get_path()

    tmp_name = timer_input(
        "\033[32m待生成的文件名称(按回车确认,ctrl+c取消):\033[m", default="res"
    )
    file_name = f"{tmp_name}.docx"
    emergency_list, critical_list, common_list = get_images(IMAGE_DIR)
    if len(common_list) + len(critical_list) + len(emergency_list) == 0:
        debug_log(
            f"未找到任何图片,结束运行",
            1,
        )
        return
    if get_template(emergency_list, critical_list, common_list, template_file_name):
        deal(emergency_list, critical_list, common_list, file_name)
        debug_log(f"请查看 \033[32m{file_name}\033[m 文件")


if __name__ == "__main__":
    debug_log("程序开始运行...")
    main()
    debug_log(f"程序运行结束！")
